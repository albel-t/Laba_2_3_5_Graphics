# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_UNDERLINE
# from docx.shared import Inches, Cm, Pt
# import glob
# import os

# def addCode(name, paragraph):
#     p = paragraph.add_run(name)
#     p.font.name = 'Consolas'
#     p.font.size = Pt(12)
#     return p

# def addTask(task, paragraph):
#     p = paragraph.add_run(task)
#     p.font.name = 'Times_New_Roman'
#     p.font.size = Pt(14)
#     return p

# def addCodeTitle(code, paragraph):
#     p = addTask(code, paragraph)
#     p.font.bold = True
#     return p

# def addTaskTitle(text, paragraph):
#     p = addTask(text, paragraph)
#     p.font.bold = True
#     return p


# def addImage(path, image, paragraph):
#     p = addTask(f'\n{image}:' , paragraph)
#     p.add_picture(path + '\\' + image, width=Inches(4.0))
#     return p


# #visual studio
# # path_to_project = input()
    
# LOG_FILE = 'loggs.txt' 

# path_to_this_document = os.path.dirname(os.path.abspath(__file__))

# path_to_template_document = "D:\\projects\\Templates\\шаблон отчета Кафедра САПР.docx"
# path_to_imgs = path_to_this_document + '\\img'
# name_laba = "laba_1"

# doc = Document(path_to_template_document)


# style = doc.styles['Normal']
# style.font.name = 'Times_New_Roman'
# style.font.size = Pt(14)
# p = doc.add_paragraph(" ")

# addTaskTitle("\nЦель работы:\n", p)
# print("Purpose of the work:")
# addTask(input() + "\n", p)

# addTaskTitle("\nЗаданиe:\n", p)
# print("Exercise:")
# addTask(input() + "\n", p)


# file_log = open(LOG_FILE, "r", encoding="utf8").readlines()
# files = []
# addTaskTitle("\nРешение:\n", p)
# for string in file_log:
#     string = string.split('|')
#     if "-init" in string[0]:
#         files.append(string[1].replace('\n', ''))
#         print("добавлени путь:" + string[1])


# for file_name in files:
#     addCodeTitle("\nИз файла ..." + file_name[-25:] + '.cpp\n', p)
#     addCode(open(file_name, "r").read(), p)
#     print("записан путь:" + file_name)


# for item in os.listdir(path_to_imgs):
#     print("добавлени путь:" + item)
#     addImage(path_to_imgs, item, p)

# doc.save(path_to_this_document + '/Oтчет_'+name_laba+'.docx') 


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches, Cm, Pt
import glob
import os

def addCode(name, paragraph):
    p = paragraph.add_run(name)
    p.font.name = 'Consolas'
    p.font.size = Pt(12)
    return p

def addTask(task, paragraph):
    p = paragraph.add_run(task)
    p.font.name = 'Times_New_Roman'
    p.font.size = Pt(14)
    return p
 
def addCodeTitle(code, paragraph):
    p = addTask(code, paragraph)
    p.font.bold = True
    return p

def addTaskTitle(text, paragraph):
    p = addTask(text, paragraph)
    p.font.bold = True
    return p

def addImage(path, image, paragraph):
    p = addTask(f'\n{image}:' , paragraph)
    p.add_picture(path + '\\' + image, width=Inches(4.0))
    return p
    
LOG_FILE = 'loggs.txt' 

path_to_this_document = os.path.dirname(os.path.abspath(__file__))

path_to_template_document = "D:\\projects\\Templates\\шаблон отчета Кафедра САПР.docx"
path_to_README = "README.md"
path_to_imgs = path_to_this_document + '\\img'
name_laba = "laba_2"

doc = Document(path_to_template_document)

style = doc.styles['Normal']
style.font.name = 'Times_New_Roman'
style.font.size = Pt(14)
p = doc.add_paragraph(" ")

file_log = open(LOG_FILE, "r", encoding="utf8").readlines()
files = []
for string in file_log:
    string = string.split('|')
    if "-init" in string[0]:
        files.append(string[1].replace('\n', ''))
        print("добавлени путь:" + string[1])


file_README = open(path_to_README, "r", encoding="utf8").read()
for paragraph_README in file_README.split("##"):
    times_README = paragraph_README.split(":")
    print(f"\n{times_README}\n")
    if len(times_README) > 1:
        addTaskTitle(f"\n{times_README[0]}\n", p)
        print(f"\n{times_README[1]}\n")
        addTask(times_README[1] + "\n", p)

addTaskTitle("\nЛистинг:\n", p)

for file_name in files:
    print("записан путь:" + file_name)
    addCodeTitle("\nИз файла ..." + file_name[-25:] + '.cpp\n', p)
    addCode(open(file_name, "r", encoding="utf8").read(), p)


for item in os.listdir(path_to_imgs):
    print("добавлени путь:" + item)
    addImage(path_to_imgs, item, p)


addTask("\n\n\n Ссылка https://github.com/albel-t/Laba_2_1_5_Graphics.git на главный репозиторий\n", p)
addTaskTitle("\nОтчет сгенерирован " + __file__.split("\\")[-1] + ": ", p)
addCode('''
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches, Cm, Pt
import glob
import os

def addCode(name, paragraph):
    p = paragraph.add_run(name)
    p.font.name = 'Consolas'
    p.font.size = Pt(12)
    return p

def addTask(task, paragraph):
    p = paragraph.add_run(task)
    p.font.name = 'Times_New_Roman'
    p.font.size = Pt(14)
    return p
 
def addCodeTitle(code, paragraph):
    p = addTask(code, paragraph)
    p.font.bold = True
    return p

def addTaskTitle(text, paragraph):
    p = addTask(text, paragraph)
    p.font.bold = True
    return p

def addImage(path, image, paragraph):
    p = addTask(f'\n{image}:' , paragraph)
    p.add_picture(path + '\\' + image, width=Inches(4.0))
    return p

#visual studio
# path_to_project = input()
    
LOG_FILE = 'loggs.txt' 

path_to_this_document = os.path.dirname(os.path.abspath(__file__))

path_to_template_document = "D:\\projects\\Templates\\шаблон отчета Кафедра САПР.docx"
path_to_README = "README.md"
path_to_imgs = path_to_this_document + '\\img'
name_laba = "laba_1"

doc = Document(path_to_template_document)

style = doc.styles['Normal']
style.font.name = 'Times_New_Roman'
style.font.size = Pt(14)
p = doc.add_paragraph(" ")

file_log = open(LOG_FILE, "r", encoding="utf8").readlines()
files = []
for string in file_log:
    string = string.split('|')
    if "-init" in string[0]:
        files.append(string[1].replace('\n', ''))
        print("добавлени путь:" + string[1])


file_README = open(path_to_README, "r", encoding="utf8").read()
for paragraph_README in file_README.split("##"):
    times_README = paragraph_README.split(":")
    print(f"\n{times_README}\n")
    if len(times_README) > 1:
        addTaskTitle(f"\n{times_README[0]}\n", p)
        print(f"\n{times_README[1]}\n")
        addTask(times_README[1] + "\n", p)


addTaskTitle("\nЛистинг:\n", p)

for file_name in files:
    print("записан путь:" + file_name)
    addCodeTitle("\nИз файла ..." + file_name[-25:] + '.cpp\n', p)
    addCode(open(file_name, "r", encoding="utf8").read(), p)

for item in os.listdir(path_to_imgs):
    print("добавлени путь:" + item)
    addImage(path_to_imgs, item, p)

doc.save(path_to_this_document + '/Oтчет_'+name_laba+'.docx')
''', p)



doc.save(path_to_this_document + '/Oтчет_'+name_laba+'.docx')